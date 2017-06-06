from tkinter import *
from backend import Database
import interface as interface
from docx import Document

d = Document()
BE = Database()

class Frontend:
    def viewCommand(self):
        interface.list1.delete(0,END)
        for row in BE.view():
            interface.list1.insert(END, row)

    def searchCommand(self):
        interface.list1.delete(0,END)
        for row in BE.search(interface.titleInput.get(),
        interface.autherInput.get(), interface.yearInput.get(),
        interface.ISBNInput.get()):
            interface.list1.insert(END,row)

    def addBook(self):
        BE.insert(interface.titleInput.get(),
        interface.autherInput.get(), interface.yearInput.get(),
        interface.ISBNInput.get())
        newEntry = (interface.titleInput.get(),
        interface.autherInput.get(), interface.yearInput.get(),
        interface.ISBNInput.get())
        interface.list1.insert(END, newEntry)

    def getSelectedRow(self,event):
        global selectedRow #makes it a global variable
        index = interface.list1.curselection()[0]
        selectedRow = interface.list1.get(index)
        interface.titleEntry.delete(0,END)
        interface.autherEntry.delete(0,END)
        interface.yearEntry.delete(0,END)
        interface.ISBNEntry.delete(0,END)
        interface.titleEntry.insert(END, selectedRow[1])
        interface.autherEntry.insert(END, selectedRow[2])
        interface.yearEntry.insert(END, selectedRow[3])
        interface.ISBNEntry.insert(END, selectedRow[4])

    def updateRow(self):
        id = selectedRow[0]
        title = interface.titleInput.get()
        author = interface.autherInput.get()
        year = interface.yearInput.get()
        isbn = interface.ISBNInput.get()
        BE.update(id,title,author,year,isbn)
        self.viewCommand()

    def deleteItem(self):
        id = selectedRow[0]
        BE.delete(id)
        self.viewCommand()

    def convertToWord(self):
        d.add_heading('Books')
        for row in BE.view():
            titleHeading = 'Title: ', str(row[1])
            AutherPara = 'Auther: ', str(row[2])
            publishedPara = 'Published in : ', str(row[3])
            ISBNPara = 'ISBN number: ', str(row[3])
            d.add_heading(titleHeading, 2)
            d.add_paragraph(AutherPara)
            d.add_paragraph(publishedPara)
            d.add_paragraph(ISBNPara)
        d.save('books.docx')
